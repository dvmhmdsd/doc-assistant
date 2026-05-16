"""Unit tests for :class:`DocxParser` (T025)."""
from __future__ import annotations

from pathlib import Path

from docx import Document

from src.parsers.docx import DocxParser


def _write_docx(path: Path) -> None:
    doc = Document()
    doc.add_heading("Agreement", level=1)
    doc.add_paragraph("Preamble.")
    doc.add_heading("Article 1", level=2)
    doc.add_paragraph("Article 1 body.")
    doc.add_heading("Article 2", level=2)
    doc.add_paragraph("Article 2 body.")
    doc.add_heading("Annex", level=1)
    doc.add_paragraph("Annex body.")
    doc.save(str(path))


def test_section_path_reflects_heading_hierarchy(tmp_path: Path) -> None:
    p = tmp_path / "sample.docx"
    _write_docx(p)

    segments = DocxParser().parse(str(p))
    by_text = {s.text.strip(): s for s in segments}

    assert by_text["Preamble."].section_path == "Agreement"
    assert by_text["Article 1 body."].section_path == "Agreement > Article 1"
    assert by_text["Article 2 body."].section_path == "Agreement > Article 2"
    assert by_text["Annex body."].section_path == "Annex"


def test_segments_have_no_page_numbers(tmp_path: Path) -> None:
    p = tmp_path / "sample.docx"
    _write_docx(p)
    segments = DocxParser().parse(str(p))
    assert all(s.page_number is None for s in segments)


def test_empty_paragraphs_are_skipped(tmp_path: Path) -> None:
    p = tmp_path / "with_blanks.docx"
    doc = Document()
    doc.add_paragraph("non-empty")
    doc.add_paragraph("")
    doc.add_paragraph("   ")  # whitespace only
    doc.add_paragraph("also non-empty")
    doc.save(str(p))

    segments = DocxParser().parse(str(p))
    assert [s.text for s in segments] == ["non-empty", "also non-empty"]


def test_table_cells_are_included(tmp_path: Path) -> None:
    p = tmp_path / "with_table.docx"
    doc = Document()
    doc.add_heading("Schedule", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Cell A"
    table.rows[0].cells[1].text = "Cell B"
    doc.save(str(p))

    segments = DocxParser().parse(str(p))
    cell_texts = {s.text for s in segments}
    assert "Cell A" in cell_texts
    assert "Cell B" in cell_texts
