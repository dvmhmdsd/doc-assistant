"""DocumentParser contract — PDF + DOCX impls must satisfy the same rules.

Generates fixture documents at test time via the same libraries the
parsers use (PyMuPDF, python-docx). Both libs are hard runtime deps in
``pyproject.toml`` — no ``try: import`` skipif.
"""
from __future__ import annotations

from pathlib import Path

import fitz  # PyMuPDF
import pytest
from docx import Document

from src.api.errors import UnsupportedMediaType
from src.parsers.factory import parser_for

_PDF_MIME = "application/pdf"
_DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def _make_pdf(path: Path, pages: list[str]) -> None:
    doc = fitz.open()
    for text in pages:
        page = doc.new_page()
        if text:
            page.insert_text((72, 72), text)
    doc.save(str(path))
    doc.close()


def _make_docx(path: Path) -> None:
    doc = Document()
    doc.add_heading("Title", level=1)
    doc.add_paragraph("Intro paragraph.")
    doc.add_heading("Section A", level=2)
    doc.add_paragraph("Body of section A.")
    doc.add_heading("Section B", level=2)
    doc.add_paragraph("Body of section B.")
    doc.save(str(path))


def test_pdf_parser_returns_ordered_segments_with_page_numbers(tmp_path: Path) -> None:
    p = tmp_path / "sample.pdf"
    _make_pdf(p, [f"page {i} text." for i in range(3)])

    parser = parser_for(str(p), _PDF_MIME)
    segments = parser.parse(str(p))

    assert segments, "parser returned no segments for a text PDF"
    assert [s.char_start for s in segments] == sorted(s.char_start for s in segments)
    assert all(s.page_number is not None for s in segments)
    assert all(s.section_path is None for s in segments)


def test_scan_only_pdf_returns_empty(tmp_path: Path) -> None:
    p = tmp_path / "blank.pdf"
    _make_pdf(p, ["", ""])
    parser = parser_for(str(p), _PDF_MIME)
    assert parser.parse(str(p)) == []


def test_docx_parser_carries_section_path_from_heading_stack(tmp_path: Path) -> None:
    p = tmp_path / "sample.docx"
    _make_docx(p)

    parser = parser_for(str(p), _DOCX_MIME)
    segments = parser.parse(str(p))

    assert segments, "docx parser returned no segments"
    section_a_body = next(
        (s for s in segments if s.text.strip() == "Body of section A."), None
    )
    assert section_a_body is not None, "missing Section A body"
    assert section_a_body.section_path == "Title > Section A"

    section_b_body = next(
        (s for s in segments if s.text.strip() == "Body of section B."), None
    )
    assert section_b_body is not None
    assert section_b_body.section_path == "Title > Section B"


def test_parser_for_rejects_unsupported_type(tmp_path: Path) -> None:
    p = tmp_path / "data.txt"
    p.write_text("plain text")

    with pytest.raises(UnsupportedMediaType):
        parser_for(str(p), "text/plain")
